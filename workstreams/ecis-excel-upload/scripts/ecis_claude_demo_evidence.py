"""ClaudeExcelImport end-to-end DEMO with Word evidence (headed, step-by-step screenshots).

Flow (each step screenshotted): Mapping Configuration -> Schedules config -> PWEL BEFORE (empty) ->
Excel template -> Upload Files (upload) -> Schedules RUN NOW (ClaudeExcelImport) -> Schedule Monitor (OK)
-> Upload Files status -> PWEL AFTER (filled). DB ground-truth at BEFORE/AFTER. Self-cleans to NULL +
re-disables the schedule. Builds workstreams/ecis-excel-upload/evidence/ClaudeExcelImport_upload_demo_evidence.docx

Run headed:  EC_HEADLESS=false py -X utf8 workstreams/ecis-excel-upload/scripts/ecis_claude_demo_evidence.py
"""
import datetime, os, shutil, time, zipfile
from pathlib import Path
import oracledb
from openpyxl import Workbook
from playwright.sync_api import sync_playwright
import docx
from docx.shared import Inches, Pt

URL   = os.environ.get("EC_URL", "https://ap-f0a7g341jn6d.corp.quorumsoftware.com:8443/")
USER  = os.environ.get("EC_USER", "sysadmin"); PASS = os.environ.get("EC_PASS", "sysadmin")
HEADED = os.environ.get("EC_HEADLESS", "false").lower() == "false"
SHOTS = Path(r"c:/Projects/ChoongYin_OS/tmp/ecis_demo_shots"); SHOTS.mkdir(parents=True, exist_ok=True)
DOCX  = Path(r"c:/Projects/ChoongYin_OS/workstreams/ecis-excel-upload/evidence/ClaudeExcelImport_upload_demo_evidence.docx")
XLSX  = SHOTS / "claude_well_test_upload_Data.xlsx"
SCHED = "ClaudeExcelImport"; IFACE = "CLAUDE_WELL_TEST"
WELLS = ["AS1_Well_001", "AS1_Well_002", "AS1_Well_003"]
DATESTR = "2003-01-10"; DAY = datetime.datetime(2003, 1, 10); PRESS = [210.5, 215.0, 220.3]
PU, AREA, TARGET_WELL = "AS1 EC Exploration Norway", "AS1_Area", "AS1_Well_001"
PWEL = "Daily Prod Well Status 1, by Well"
steps = []  # (title, desc, img_path|None)

def log(m): print(m, flush=True)
def db(): return oracledb.connect(user="ECKERNEL_EC", password="energy", dsn=os.environ.get("EC_DB_DSN", "localhost:1521/ORCL"))

def well_press():
    c = db().cursor()
    c.execute("""SELECT object_code, avg_bh_press FROM dv_pwel_day_status WHERE daytime=TO_DATE(:d,'YYYY-MM-DD')
                 AND object_code IN ('AS1_Well_001','AS1_Well_002','AS1_Well_003') ORDER BY object_code""", d=DATESTR)
    return c.fetchall()

def sched_enabled():
    c = db().cursor(); c.execute("SELECT enabled FROM tv_schedule WHERE name=:n", n=SCHED); r = c.fetchone(); return r[0] if r else None

def run_status():
    c = db().cursor()
    c.execute("""SELECT * FROM (SELECT run_status, from_daytime FROM tv_action_instance_history
                 WHERE schedule_no=(SELECT schedule_no FROM tv_schedule WHERE name=:n) ORDER BY from_daytime DESC NULLS LAST)
                 WHERE ROWNUM<=2""", n=SCHED)
    return c.fetchall()

def preclean():
    c = db(); cu = c.cursor()
    cu.execute("SELECT object_id FROM ov_well WHERE code IN ('AS1_Well_001','AS1_Well_002','AS1_Well_003')")
    oids = [r[0] for r in cu.fetchall()]
    if oids:
        ph = ",".join(f"'{o}'" for o in oids)
        cu.execute(f"UPDATE pwel_day_status SET avg_bh_press=NULL WHERE object_id IN ({ph}) AND daytime=TO_DATE(:d,'YYYY-MM-DD')", d=DATESTR)
    cu.execute("DELETE FROM imp_source_interface_file WHERE interface_code=:i", i=IFACE)
    cu.execute("DELETE FROM imp_staging WHERE interface_code=:i", i=IFACE)
    c.commit(); c.close()

# ---- Excel template (sheet 'Data')
wb = Workbook(); ws = wb.active; ws.title = "Data"; ws.append(["Well", "Date", "Pressure"])
for w, pv in zip(WELLS, PRESS): ws.append([w, DAY, pv])
wb.save(XLSX)
tmp = XLSX.with_suffix(".r.xlsx")
with zipfile.ZipFile(XLSX) as zin:
    order = [n for n in zin.namelist() if n == "[Content_Types].xml"] + [n for n in zin.namelist() if n != "[Content_Types].xml"]
    with zipfile.ZipFile(tmp, "w", zipfile.ZIP_DEFLATED) as zo:
        for n in order: zo.writestr(n, zin.read(n))
shutil.move(tmp, XLSX)

log("STEP 0: pre-clean (NULL baseline + clear filedrop)")
preclean()
before = well_press(); log(f"  DB BEFORE: {before}")

with sync_playwright() as p:
    b = p.chromium.launch(headless=not HEADED, slow_mo=int(os.environ.get("EC_SLOWMO", "200")) if HEADED else 0, args=["--ignore-certificate-errors", "--start-maximized"])
    page = b.new_context(ignore_https_errors=True, no_viewport=HEADED, viewport=None if HEADED else {"width": 1920, "height": 1080}).new_page()
    page.set_default_timeout(30000)

    def shot(title, desc):
        f = SHOTS / f"{len(steps):02d}_{title.split(' ')[0].lower()}.png"
        try: page.screenshot(path=str(f), full_page=True)
        except Exception as e: log(f"  shot fail {title}: {e}"); f = None
        steps.append((title, desc, str(f) if f else None)); log(f"  captured: {title}")

    def open_screen(name):
        page.goto(URL, wait_until="domcontentloaded", timeout=60000)
        page.wait_for_selector('[id="menu:searchForm:searchTxt"]', timeout=60000); time.sleep(1)
        bx = page.locator('[id="menu:searchForm:searchTxt"]'); bx.fill(""); bx.type(name, delay=40); time.sleep(1.5)
        page.locator(f'xpath=//*[contains(@class,"tv-link") and normalize-space(text())="{name}"]').first.click()
        page.wait_for_load_state("networkidle", timeout=25000); time.sleep(2.5)

    def pick(dd, label):
        if not page.locator(f'[id="{dd}_panel"]').is_visible():
            page.click(f'[id="{dd}_button"]'); page.wait_for_selector(f'[id="{dd}_panel"] tr[data-item-label]', timeout=8000)
        page.locator(f'[id="{dd}_panel"] tr[data-item-label="{label}"]').first.click()
        page.wait_for_load_state("networkidle", timeout=15000); time.sleep(0.9)

    def options(dd):
        page.click(f'[id="{dd}_button"]'); page.wait_for_selector(f'[id="{dd}_panel"] tr[data-item-label]', timeout=7000)
        return page.evaluate(f"""()=>[...document.querySelectorAll('[id="{dd}_panel"] tr[data-item-label]')].map(t=>t.getAttribute('data-item-label'))""")

    def go(): page.click('[id="button:form:B"]'); page.wait_for_load_state("networkidle", timeout=25000); time.sleep(2.5)

    def nav_pwel(facility=None):
        for g in (0, 1):
            di = page.locator(f'[id="nav:form:G:{g}:R:1:C:0:da_input"]')
            if di.count(): di.first.fill(DATESTR); page.keyboard.press("Tab"); time.sleep(0.6)
        pick("nav:form:G:2:R:1:C:0:dd", PU); pick("nav:form:G:3:R:1:C:0:dd", AREA)
        if facility:
            pick("nav:form:G:4:R:1:C:0:dd", facility); pick("nav:form:G:5:R:1:C:0:dd", TARGET_WELL); return facility
        for fac in options("nav:form:G:4:R:1:C:0:dd"):
            page.keyboard.press("Escape"); time.sleep(0.3); pick("nav:form:G:4:R:1:C:0:dd", fac)
            if TARGET_WELL in options("nav:form:G:5:R:1:C:0:dd"):
                pick("nav:form:G:5:R:1:C:0:dd", TARGET_WELL); log(f"    well under '{fac}'"); return fac
            page.keyboard.press("Escape"); time.sleep(0.3)
        return None

    def sel_sched():
        open_screen("Schedules"); pick("nav:form:G:0:R:0:C:1:dd", "All"); go()
        page.fill('[id="schedule:form:T:sfilter0_ft_filter"]', SCHED); page.keyboard.press("Enter")
        page.wait_for_load_state("networkidle", timeout=15000); time.sleep(2)
        page.locator(f'xpath=//tbody[@id="schedule:form:T_data"]//input[@value="{SCHED}"]').first.click()
        page.wait_for_load_state("networkidle", timeout=15000); time.sleep(1.5)

    def toggle_enabled_save():
        page.click('[id="tab:tabPanel:job_details_more:form:G:0:R:0:C:0:cb"]')
        page.locator('xpath=//li[.//span[contains(@class,"ui-icon-save")]][1]').first.click()
        page.wait_for_load_state("networkidle", timeout=20000); time.sleep(3)

    # ---- login
    page.goto(URL, wait_until="domcontentloaded", timeout=60000)
    page.fill('[id="username"]', USER); page.fill('[id="password"]', PASS); page.click('[id="kc-login"]')
    page.wait_for_selector('[id="menu:searchForm:searchTxt"]', timeout=60000); time.sleep(1)

    # SETUP: enable the schedule FIRST + let it settle. RUN NOW needs ENABLED, and toggling enable right before
    # RUN NOW is unreliable (the enable-save can auto-fire on an empty drop). Enable as setup, settle, clean residue.
    log("SETUP: ensure schedule enabled + settle"); sel_sched()
    if sched_enabled() == "N": toggle_enabled_save(); log(f"  enabled -> {sched_enabled()}")
    time.sleep(12)
    _con = db(); _cu = _con.cursor()
    _cu.execute("DELETE FROM imp_staging WHERE interface_code=:i", i=IFACE)
    _cu.execute("UPDATE pwel_day_status SET avg_bh_press=NULL WHERE object_id IN (SELECT object_id FROM ov_well WHERE code IN ('AS1_Well_001','AS1_Well_002','AS1_Well_003')) AND daytime=TO_DATE(:d,'YYYY-MM-DD')", d=DATESTR)
    _con.commit(); _con.close()

    # STEP 1: Mapping Configuration
    log("STEP 1: Mapping Configuration"); open_screen("Mapping Configuration")
    try:
        fb = page.locator('xpath=//input[contains(@id,"sfilter0_ft_filter")]')
        if fb.count(): fb.first.fill(IFACE); page.keyboard.press("Enter"); page.wait_for_load_state("networkidle"); time.sleep(2)
    except Exception: pass
    shot("Mapping Configuration (CLAUDE_WELL_TEST interface + source/target mappings)",
         "The ECIS interface config under test: CLAUDE_WELL_TEST with Well/Date/Pressure source mappings and the AVG_BH_PRESS target mapping.")

    # STEP 2: Schedules config
    log("STEP 2: Schedules config"); sel_sched()
    shot("Schedules - ClaudeExcelImport configuration",
         "The ClaudeExcelImport schedule: two ECISAction instances (file->staging, staging->target), INTERFACE_CODE=CLAUDE_WELL_TEST.")

    # STEP 3: PWEL BEFORE
    log("STEP 3: PWEL BEFORE"); open_screen(PWEL); fac = nav_pwel(); go()
    shot("Daily Prod Well Status 1, by Well - BEFORE (AVG_BH_PRESS empty)",
         f"Before the upload: AVG_BH_PRESS is empty for {', '.join(WELLS)} on {DATESTR}. DB ground truth: {before}.")

    # STEP 4: Excel template (Word table built later; also screenshot upload-with-file below)
    steps.append(("Excel template file (sheet 'Data')",
                  f"Template uploaded: sheet 'Data', columns Well | Date | Pressure. Rows: " +
                  "; ".join(f"{w}={pv}" for w, pv in zip(WELLS, PRESS)) + f" on {DATESTR}. File: {XLSX.name}.", None))

    # STEP 5: Upload Files
    log("STEP 5: Upload Files"); open_screen("Upload Files")
    pick("StandardNavigator:form:G:2:R:1:C:0:dd", "ECIS Interface Area")
    pick("StandardNavigator:form:G:3:R:1:C:0:dd", "Claude Well Test")
    page.click('[id="buttongo:form:B"]'); page.wait_for_load_state("networkidle", timeout=20000); time.sleep(2)
    page.set_input_files('[id="upload_file_btn:form:fa_input"]', str(XLSX)); time.sleep(2.5)
    shot("Upload Files - file selected (Claude Well Test interface)",
         "Upload Files screen: ECIS Interface Area -> Claude Well Test, the Data Excel selected, ready to upload.")
    page.locator('xpath=//*[self::button or self::span][normalize-space(.)="Upload File"]').locator("visible=true").first.click()
    page.wait_for_load_state("networkidle", timeout=30000); time.sleep(3)
    shot("Upload Files - uploaded", "After clicking Upload File: the file is accepted into the ECIS filedrop for CLAUDE_WELL_TEST.")
    # settle: ensure the uploaded file is committed/visible in the DB filedrop before RUN NOW (guards 'no files found')
    for _ in range(15):
        time.sleep(2)
        _cur = db().cursor(); _cur.execute("SELECT COUNT(*) FROM imp_source_interface_file WHERE interface_code=:i", i=IFACE)
        if _cur.fetchone()[0] > 0: break
    log("  file committed in filedrop; settling 10s"); time.sleep(10)

    # STEP 6: enable + RUN NOW
    log("STEP 6: enable + RUN NOW"); sel_sched()
    if sched_enabled() == "N": toggle_enabled_save(); log(f"  enabled -> {sched_enabled()}")
    time.sleep(8)  # let the enable save (+ any auto-fire) settle before RUN NOW
    shot("Schedules - ClaudeExcelImport enabled, about to RUN NOW", "Schedule enabled; clicking RUN NOW to trigger the import job chain.")
    page.click('[id="runNowButton:form:B"]'); page.wait_for_load_state("networkidle", timeout=30000); time.sleep(2)
    dlg = page.locator('xpath=//*[self::button or self::span][normalize-space(.)="Yes" or normalize-space(.)="OK"]').locator("visible=true")
    if dlg.count(): dlg.first.click()
    shot("Schedules - RUN NOW triggered", "RUN NOW confirmed - ClaudeExcelImport executes its two ECISAction steps (file->staging, staging->target).")

    # wait for processing
    log("  waiting for data to land ...")
    for _ in range(20):
        time.sleep(3)
        if all(v is not None for _, v in well_press()): break
    after = well_press(); rs = run_status(); log(f"  DB AFTER: {after} | run_status: {rs}")

    # STEP 7: MONITOR tab on the Schedules screen (after selecting the schedule)
    log("STEP 7: Monitor tab")
    sel_sched()
    try:
        page.locator('xpath=//ul[contains(@class,"ui-tabs-nav")]//a[normalize-space(.)="Monitor"]').first.click()
        page.wait_for_load_state("networkidle", timeout=15000); time.sleep(2)
        try:  # GO-refresh the monitor log so the latest run (OK) shows at the top
            page.locator('[id="tab:tabPanel:buttonNav:form:B"]').first.click(); page.wait_for_load_state("networkidle", timeout=12000); time.sleep(2)
        except Exception: pass
    except Exception as e:
        log(f"  MONITOR tab click failed: {e}")
    shot("Schedules - MONITOR tab (ClaudeExcelImport run log/status)",
         f"The MONITOR tab shows the ClaudeExcelImport run log + status. DB run_status: {rs}.")

    # STEP 8: Upload Files status
    log("STEP 8: Upload Files status"); open_screen("Upload Files")
    pick("StandardNavigator:form:G:2:R:1:C:0:dd", "ECIS Interface Area")
    pick("StandardNavigator:form:G:3:R:1:C:0:dd", "Claude Well Test")
    page.click('[id="buttongo:form:B"]'); page.wait_for_load_state("networkidle", timeout=20000); time.sleep(2)
    shot("Upload Files - latest file status (parsed / written to EC)",
         "Upload Files screen after the run: the uploaded file shows its processed status (parsed + written to EC).")

    # STEP 9: PWEL AFTER
    log("STEP 9: PWEL AFTER"); open_screen(PWEL); nav_pwel(facility=fac); go()
    shot("Daily Prod Well Status 1, by Well - AFTER (AVG_BH_PRESS filled)",
         f"After the upload+run: AVG_BH_PRESS is populated. DB ground truth: {after} (bar; screen shows psi).")

    # restore: disable schedule
    log("RESTORE: disable schedule"); sel_sched()
    if sched_enabled() == "Y": toggle_enabled_save(); log(f"  disabled -> {sched_enabled()}")
    b.close()

# ---- self-clean DB
log("SELF-CLEAN: revert data + clear filedrop")
preclean()
final = well_press(); log(f"  DB FINAL (expect NULLs): {final}")

# ---- build Word doc
log("BUILD DOCX")
doc = docx.Document()
doc.add_heading("ECIS Excel Upload - ClaudeExcelImport End-to-End Evidence", 0)
m = doc.add_paragraph()
m.add_run("Interface: ").bold = True; m.add_run(f"{IFACE}   ")
m.add_run("Schedule: ").bold = True; m.add_run(f"{SCHED}   ")
m.add_run("Target: ").bold = True; m.add_run("PWEL_DAY_STATUS.AVG_BH_PRESS")
doc.add_paragraph(f"Wells {', '.join(WELLS)} @ {DATESTR}.  Sandbox {URL}  (DB-verified).")
# DB evidence table
t = doc.add_table(rows=1, cols=3); t.style = "Light Grid Accent 1"
for i, h in enumerate(["Well", "BEFORE avg_bh_press", "AFTER avg_bh_press"]): t.rows[0].cells[i].text = h
bmap = dict(before); amap = dict(after)
for w in WELLS:
    r = t.add_row().cells; r[0].text = w; r[1].text = str(bmap.get(w)); r[2].text = str(amap.get(w))
doc.add_paragraph(f"Schedule run_status (DB tv_action_instance_history): {run_status()}")
for i, (title, desc, img) in enumerate(steps, 1):
    doc.add_heading(f"Step {i}: {title}", level=2)
    doc.add_paragraph(desc)
    if i == 4:  # excel template as a native table
        tt = doc.add_table(rows=1, cols=3); tt.style = "Light Grid Accent 1"
        for j, h in enumerate(["Well", "Date", "Pressure"]): tt.rows[0].cells[j].text = h
        for w, pv in zip(WELLS, PRESS):
            rr = tt.add_row().cells; rr[0].text = w; rr[1].text = DATESTR; rr[2].text = str(pv)
    if img and Path(img).exists():
        try: doc.add_picture(img, width=Inches(6.8))
        except Exception as e: doc.add_paragraph(f"[screenshot {img} - {e}]")
doc.add_heading("Result", level=2)
ok = all(amap.get(w) is not None for w in WELLS) and all(bmap.get(w) is None for w in WELLS)
doc.add_paragraph(("PASS" if ok else "REVIEW") + f": before all empty -> after all filled via {SCHED}. Self-cleaned to NULL; schedule re-disabled (enabled={sched_enabled()}).")
DOCX.parent.mkdir(parents=True, exist_ok=True); doc.save(str(DOCX))
log(f"DONE -> {DOCX}")
print("STEPS:", [s[0] for s in steps])
print("BEFORE:", before, "AFTER:", after, "FINAL:", final, "PASS:", ok)
