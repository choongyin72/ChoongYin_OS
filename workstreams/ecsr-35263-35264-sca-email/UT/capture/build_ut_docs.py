# -*- coding: utf-8 -*-
"""
Assemble the two ECSR-35264 UT docs (PLUTO + SCA) in the 11-screen NOPTA format
(model: workstreams/ecsr-35329-35330-nopta-email/UT/UT_ECPR-31089.docx).

Inputs (same folder):  ./shots/<variant>/NN_*.png  (from capture_ut_screens.py)
                       ./content.json               (from fetch_message_content.py)
Output:                ../UT_ECSR-35264__<DEF>.docx
Usage:                 py build_ut_docs.py
"""
import os, json
from docx import Document
from docx.shared import Inches, Pt

HERE  = os.path.dirname(__file__)
SHOTS = os.path.join(HERE, "shots")
OUTUT = os.path.abspath(os.path.join(HERE, ".."))
prev  = json.load(open(os.path.join(HERE, "content.json"), encoding="utf-8"))

VAR = {
    "pluto": {"label": "Pluto", "def": "R_BLP_DAILY_PROD_ALLOC_PLUTO", "set": "R_BLP_DAILY_PROD_ALLOC_PLU",
              "report": "Burrup LNG Park Daily Production Report (Pluto)",
              "subj_tmpl": "Burrup LNG Park Daily Production Report {production_day}"},
    "sca":   {"label": "Scarborough", "def": "R_BLP_DAILY_PROD_ALLOC_SCA", "set": "R_BLP_DAILY_PROD_ALLOC_SCA",
              "report": "Burrup LNG Park Daily Production Report (Scarborough)",
              "subj_tmpl": "Burrup LNG Park Daily Production Report (Scarborough) {production_day}"},
}


def sections(v):
    c = VAR[v]; D = c["def"]; S = c["set"]; L = c["label"]; p = prev[v]
    mno = p["msg_no"] if p else "-"
    return [
        ("01_msgtype", "1.  Maintain Message Type screen",
         f"Dedicated message type {D} for the {c['report']}. Internal Format = Free Text, Direction = "
         f"Outbound -- a standalone definition separate from the shared R_BLP_DAILY_PROD_ALLOC type."),
        ("02_msgformat", "2.  Message Format screen",
         f"TEXT is the Default External Format (Plain text) for message type {D}."),
        ("03_freetext", "3.  Freetext Message Template screen",
         f"Freetext body template for {D} -- subject '{c['subj_tmpl']}'. The {{production_day}} placeholder "
         f"is resolved at send time; the body carries the {L} production-date narrative."),
        ("04_contactgroupset", "4.  Maintain Contact Group Set screen",
         f"Contact Group Set {S} -- the dedicated actor set for the {L} report, cloned from and kept "
         f"separate from the shared R_BLP_DAILY_PROD_ALLOC set."),
        ("05_actormaint", "5.  Actor Maintenance screen",
         f"Message contacts under the {L} set: the sender (From, SMTP) plus the internal/external report "
         f"recipients, all SMTP."),
        ("06_distlist", "6.  Distribution List screen",
         f"Distribution List {S} -- the To / From / Cc SMTP recipients for the {L} report; a distribution "
         f"separate from the shared set so {L} mail is addressed independently."),
        ("07_msgdistribution", "7.  Message Distribution screen",
         f"Message Distribution for {D}: Format = TEXT, linking the message definition to the {L} report "
         f"and the {S} distribution set."),
        ("08_reportadmin", "8.  Report Administration screen",
         f"{c['report']} in Report Administration -- the runnable report whose SEND queues an outgoing "
         f"message to the configured {L} recipients."),
        ("09_schedules", "9.  Schedules screen",
         "ZWP_SEND_BPM_NOTIFICATIONS (BPM Send Notifications) -- the scheduled business action that renders "
         "subject/body from the template (updateMsgOutFromMsgTemplate) and despatches MHM messages. Shared by all reports."),
        ("10_outgoing", "10.  Outgoing Messages screen",
         f"Generated outgoing message {mno}: Message Type {D}, subject '{p['subject'] if p else ''}', addressed "
         f"to the {L} SMTP recipients. Status = ERROR is the expected SMTP-not-configured state on COPSDEV (mail "
         f"is not despatched); the message itself is built correctly with the {L} type, subject and recipients."),
    ]


def build(v):
    c = VAR[v]; p = prev[v]; doc = Document()
    t = doc.add_paragraph("Unit Test Evidence - ECSR-35264"); t.style = doc.styles["Title"]
    doc.add_paragraph(f"Split the shared Burrup LNG Park Daily Production email configuration into a dedicated "
                      f"set for the {c['label']} report ({c['def']}).")
    doc.add_paragraph(f"Jira (UAT): ECSR-35264    Report: {c['def']}    Distribution/Contact set: {c['set']}    "
                      f"Environment: COPSDEV (= plutodev)")
    for fn, head, cap in sections(v):
        doc.add_heading(head, level=2)
        doc.add_paragraph(cap)
        doc.add_picture(os.path.join(SHOTS, v, fn + ".png"), width=Inches(6.6))
    doc.add_heading("11.  Preview generated outgoing message", level=2)
    doc.add_paragraph(f"Generated outgoing message {p['msg_no']} for {c['def']} -- the rendered email below "
                      f"confirms the {c['label']}-specific subject and body produced by this configuration.")
    doc.add_picture(os.path.join(SHOTS, v, "10_outgoing.png"), width=Inches(6.6))
    doc.add_paragraph(f"Rendered subject: '{p['subject']}'")
    doc.add_paragraph("Rendered message body (MESSAGE_OUT.MESSAGE_DRAFT):")
    pb = doc.add_paragraph(p["body"])
    for r in pb.runs:
        r.font.size = Pt(9)
    out = os.path.join(OUTUT, f"UT_ECSR-35264__{c['def']}.docx")
    doc.save(out); print("WROTE", out)


for v in ("pluto", "sca"):
    build(v)
print("DONE")
