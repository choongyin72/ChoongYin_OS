"""Insert screenshot blocks into the evidence doc AT a section position (mid-doc), before a
target heading. Reusable: pass the target heading prefix + the blocks to insert before it.
Each block = (heading3, caption, image_path). Preserves all existing content. No commit."""
import os, sys
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

DOC = r"c:/Projects/ChoongYin_OS/workstreams/production-stability/sql-scripts/Issue1052_Evidence_COPS_DEV.docx"

def find_heading(doc, prefix):
    for p in doc.paragraphs:
        if p.style.name.startswith('Heading') and p.text.strip().startswith(prefix):
            return p
    return None

def insert_blocks_before(target, intro_heading, intro_text, blocks):
    # intro heading (level 2) + intro paragraph
    h = target.insert_paragraph_before(intro_heading, style='Heading 2')
    ip = target.insert_paragraph_before(intro_text); ip.runs[0].font.size = Pt(9)
    for title, cap, path in blocks:
        target.insert_paragraph_before(title, style='Heading 3')
        c = target.insert_paragraph_before(cap)
        c.runs[0].font.size = Pt(8.5); c.runs[0].font.color.rgb = RGBColor(0x60,0x60,0x60)
        pic = target.insert_paragraph_before()
        pic.add_run().add_picture(path, width=Inches(6.8))
        pic.alignment = WD_ALIGN_PARAGRAPH.CENTER
        target.insert_paragraph_before('')

SS = r"c:/Projects/ChoongYin_OS/workstreams/master-plan/ec-automation/results/sum_evidence"

doc = Document(DOC)
target = find_heading(doc, '11. Sum')   # insert frozen screen evidence at end of §10, before §11
assert target is not None, "could not find §11 heading"

blocks = [
 ('10.4.1  Stream Gas Component Analysis (Analysis) — DENSITY & GCV frozen (1152/1153)',
  'Validation Overview - Pluto Scarborough > Stream Gas Component Analysis (Analysis) - PHD '
  'Validations, 2025-12-13. Message filtered to "same as previous day": 12 frozen WARNINGs '
  '(DENSITY + GCV across 6 streams) surfacing on screen.',
  os.path.join(SS, 'frozen_analysis_2025-12-13.png')),
 ('10.4.2  Daily Stream Water Status — Oil-in-Water frozen (1154)',
  'Daily Stream Water Status - PHD Validations, 2026-05-24, filtered to "same as previous day": '
  'the Oil-in-Water frozen WARNING (rule 1154) on screen.',
  os.path.join(SS, 'frozen_water_2026-05-24.png')),
]
insert_blocks_before(target, '10.4  Frozen — EC Web Screen Evidence (Validation Overview)',
    'Headless RF capture on COPS DEV. The frozen check groups are run on the Validation Overview '
    'screen; the result grid is filtered on the Message column to isolate the frozen WARNINGs.',
    blocks)
doc.save(DOC)
print("Inserted 10.4 (frozen screen evidence) before §11.")
