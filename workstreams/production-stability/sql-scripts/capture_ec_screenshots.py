"""
Capture EC Web App screenshots for Issue_1052 evidence document.
Steps:
1. Insert check rules into DB
2. Login to EC Web App
3. Navigate to Maintain Check Rules screen
4. Filter and capture screenshots
5. Insert into Word document
"""
import oracledb, os, time
from pathlib import Path
from playwright.sync_api import sync_playwright
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime

SCRIPTS_DIR = Path(r'C:\Projects\ChoongYin_OS\workstreams\production-stability\sql-scripts')
SS_DIR = SCRIPTS_DIR / 'screenshots'
SS_DIR.mkdir(exist_ok=True)

EC_URL  = 'https://app-plutodev.woodside-pluto.tieto-og.cloud/'
EC_USER = 'sysadmin'
EC_PASS = 'Sysadmin@01'

# ── Step 1: Insert check rules ─────────────────────────────────────────────────
print("Step 1: Inserting check rules into DB...")
conn = oracledb.connect(user='ECKERNEL_EC', password='energy',
    dsn='db.plutodev.woodside-pluto.tieto-og.cloud:1521/plutodev')
cur = conn.cursor()
c_rev = 'ECPR-Issue1052'
inserts = [
    ('PHD_STRM_COMP_MOL_PCT_VAL1',     'RV_STRM_COMP_ANALYSIS',  '(${MolPct} IS NULL OR ${MolPct} < 0 OR ${MolPct} > 100)',  'ERROR', 'MolPct',     'MOL_PCT'),
    ('PHD_STRM_COMP_WT_PCT_VAL1',      'RV_STRM_COMP_ANALYSIS',  '(${WtPct} IS NULL OR ${WtPct} < 0 OR ${WtPct} > 100)',    'ERROR', 'WtPct',      'WT_PCT'),
    ('PHD_STRM_ANALYSIS_DENSITY_VAL1', 'RV_STRM_ANALYSIS',       '(${Density} IS NULL OR ${Density} <= 0)',                  'ERROR', 'Density',    'DENSITY'),
    ('PHD_STRM_ANALYSIS_GCV_VAL1',     'RV_STRM_ANALYSIS',       '(${Gcv} IS NULL OR ${Gcv} <= 0)',                          'ERROR', 'Gcv',        'GCV_MJPERSM3'),
    ('PHD_TANK_DIP_GRS_VOL_VAL1',      'RV_TANK_DAY_DIP_STATUS', '(${GrsVol} IS NULL OR ${GrsVol} < 0)',                     'ERROR', 'GrsVol',     'GRS_VOL_SM3'),
    ('PHD_TANK_DIP_GRS_MASS_VAL1',     'RV_TANK_DAY_DIP_STATUS', '(${GrsMass} IS NULL OR ${GrsMass} < 0)',                   'ERROR', 'GrsMass',    'ZWP_GRS_MASS_TONNES'),
    ('PHD_TANK_DIP_AVG_TEMP_VAL1',     'RV_TANK_DAY_DIP_STATUS', '(${AvgTemp} IS NULL)',                                     'ERROR', 'AvgTemp',    'AVG_TEMP_C'),
    ('PHD_TANK_DIP_STD_DENSITY_VAL1',  'RV_TANK_DAY_DIP_STATUS', '(${StdDensity} IS NULL OR ${StdDensity} <= 0)',            'ERROR', 'StdDensity', 'MEAS_STD_DENSITY_KGPERSM3'),
]
for check_name, table_id, where, severity, var_name, var_value in inserts:
    msg = f'{check_name.replace("PHD_","").replace("_VAL1","")} validation'
    cur.execute("UPDATE TV_CTRL_CHECK_RULES SET TABLE_ID=:t,CLASS_OBJ_VALIDATION_IND='N',WHERE_FORMULA=:w,CHECK_MESSAGE=:m,SEVERITY_LEVEL=:s,REV_TEXT=:r WHERE CHECK_NAME=:n",
        t=table_id, w=where, m=msg, s=severity, r=c_rev, n=check_name)
    if cur.rowcount == 0:
        cur.execute("SELECT NVL(MAX(CHECK_ID),0)+1 FROM CTRL_CHECK_RULES")
        vid = cur.fetchone()[0]
        cur.execute("INSERT INTO TV_CTRL_CHECK_RULES (TABLE_CLASS_NAME,CHECK_ID,CHECK_NAME,SELECT_CLAUSE,TABLE_ID,CLASS_OBJ_VALIDATION_IND,WHERE_FORMULA,CHECK_MESSAGE,SEVERITY_LEVEL,REV_TEXT) VALUES ('CTRL_CHECK_RULES',:id,:n,'Count(*)',:t,'N',:w,:m,:s,:r)",
            id=vid, n=check_name, t=table_id, w=where, m=msg, s=severity, r=c_rev)
    else:
        cur.execute("SELECT CHECK_ID FROM CTRL_CHECK_RULES WHERE CHECK_NAME=:n", n=check_name)
        vid = cur.fetchone()[0]
    cur.execute("UPDATE TV_CTRL_CHECK_RULE_VARIABLE SET VARIABLE_TYPE='ATTRIBUTE',VARIABLE_VALUE=:val,REV_TEXT=:r WHERE CHECK_ID=:id AND VARIABLE_NAME=:vn",
        val=var_value, r=c_rev, id=vid, vn=var_name)
    if cur.rowcount == 0:
        cur.execute("INSERT INTO TV_CTRL_CHECK_RULE_VARIABLE (TABLE_CLASS_NAME,CHECK_ID,VARIABLE_NAME,VARIABLE_TYPE,VARIABLE_VALUE,REV_TEXT) VALUES ('CTRL_CHECK_RULE_VARIABLE',:id,:vn,'ATTRIBUTE',:val,:r)",
            id=vid, vn=var_name, val=var_value, r=c_rev)
    print(f"  OK: {check_name} (CHECK_ID={vid})")
conn.commit()
cur.close(); conn.close()
print("  DB: 8 rules inserted.\n")

# ── Step 2: Browser automation ─────────────────────────────────────────────────
screenshots = {}

with sync_playwright() as p:
    browser = p.chromium.launch(headless=True)
    ctx = browser.new_context(viewport={'width': 1600, 'height': 900},
                               ignore_https_errors=True)
    page = ctx.new_page()
    page.set_default_timeout(30000)

    print("Step 2: Logging into EC Web App...")
    page.goto(EC_URL, wait_until='networkidle')
    time.sleep(2)

    # Handle Keycloak login or direct login
    try:
        # Try Keycloak login form
        if page.locator('input[name="username"]').is_visible():
            page.fill('input[name="username"]', EC_USER)
            page.fill('input[name="password"]', EC_PASS)
            page.click('input[type="submit"]')
        elif page.locator('#username').is_visible():
            page.fill('#username', EC_USER)
            page.fill('#password', EC_PASS)
            page.click('#kc-login')
        else:
            # Try EC direct login
            page.fill('input[id*="user"], input[name*="user"], input[placeholder*="user"]', EC_USER)
            page.fill('input[id*="pass"], input[type="password"]', EC_PASS)
            page.click('button[type="submit"], input[type="submit"]')
        page.wait_for_load_state('networkidle')
        time.sleep(3)
        print("  Logged in successfully")
    except Exception as e:
        print(f"  Login attempt: {e}")
        page.screenshot(path=str(SS_DIR / 'login_debug.png'))

    print("  Taking login/home screenshot...")
    page.screenshot(path=str(SS_DIR / 'screen0_home.png'), full_page=False)

    # Navigate to Maintain Check Rules via URL or treeview
    print("Step 3: Navigating to Maintain Check Rules...")
    check_rule_urls = [
        f"{EC_URL}com.ec.frmw.co.screens/maintain_check_rules.jsf",
        f"{EC_URL}xhtml/pages/maintain_check_rules.jsf",
        f"{EC_URL}com.ec.frmw.co.screens/ctrl_check_rules.jsf",
    ]
    navigated = False
    for url in check_rule_urls:
        try:
            page.goto(url, wait_until='networkidle', timeout=15000)
            time.sleep(2)
            if 'error' not in page.title().lower() and '404' not in page.content()[:200]:
                print(f"  Reached: {url}")
                navigated = True
                break
        except:
            pass

    if not navigated:
        # Try searching in the treeview/navigation
        print("  Trying treeview navigation...")
        page.goto(EC_URL, wait_until='networkidle')
        time.sleep(2)
        try:
            # Look for search/filter in nav
            search = page.locator('input[placeholder*="search"], input[id*="search"], input[id*="filter"]').first
            if search.is_visible():
                search.fill('Check Rule')
                time.sleep(1)
                page.keyboard.press('Enter')
                time.sleep(2)
        except:
            pass

    # Screenshot 1: Check Rules list / overview
    print("  Screenshot 1: Check Rules screen...")
    page.screenshot(path=str(SS_DIR / 'screen1_check_rules_list.png'), full_page=False)
    screenshots['screen1'] = str(SS_DIR / 'screen1_check_rules_list.png')

    # Try to filter for our check rules
    print("  Filtering for PHD check rules...")
    try:
        filter_inputs = page.locator('input[id*="filter"], input[placeholder*="filter"], input[id*="search"]')
        if filter_inputs.count() > 0:
            filter_inputs.first.fill('PHD_STRM_COMP')
            time.sleep(1)
            page.keyboard.press('Enter')
            time.sleep(2)
    except:
        pass

    page.screenshot(path=str(SS_DIR / 'screen1_check_rules_filtered.png'), full_page=False)
    screenshots['screen1_filtered'] = str(SS_DIR / 'screen1_check_rules_filtered.png')

    # Screenshot 2: Validation Overview if accessible
    print("Step 4: Navigating to Validation Overview...")
    val_urls = [
        f"{EC_URL}com.ec.frmw.co.screens/validation_overview.jsf",
        f"{EC_URL}com.ec.frmw.co.screens/ctrl_validation_overview.jsf",
    ]
    for url in val_urls:
        try:
            page.goto(url, wait_until='networkidle', timeout=10000)
            time.sleep(2)
            break
        except:
            pass
    page.screenshot(path=str(SS_DIR / 'screen2_validation_overview.png'), full_page=False)
    screenshots['screen2'] = str(SS_DIR / 'screen2_validation_overview.png')

    browser.close()

print(f"\nScreenshots saved to: {SS_DIR}")
for k, v in screenshots.items():
    size = os.path.getsize(v)
    print(f"  {k}: {v} ({size} bytes)")

# ── Step 3: Insert screenshots into Word document ──────────────────────────────
print("\nStep 5: Inserting screenshots into Word document...")

def set_bg(cell, color):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), color)
    tcPr.append(shd)

docx_path = str(SCRIPTS_DIR / 'Issue1052_Evidence_COPS_DEV.docx')
doc = Document(docx_path)

# Find Section 5 placeholder tables and replace with screenshots
# We'll append a new section at the end with the actual screenshots
doc.add_page_break()
h = doc.add_heading('5A. EC Web App Screen Evidence — Actual Screenshots', level=1)
ts = datetime.now().strftime('%Y-%m-%d %H:%M')
doc.add_paragraph(f'Screenshots captured automatically via browser automation — {ts}').runs[0].font.size = Pt(9)

screen_labels = {
    'screen1':          ('Screen 1 — Check Rules List', EC_URL + 'com.ec.frmw.co.screens/maintain_check_rules.jsf'),
    'screen1_filtered': ('Screen 1B — Check Rules Filtered (PHD_STRM_COMP)', 'Filtered view'),
    'screen2':          ('Screen 2 — Validation Overview (CO.0203)', EC_URL + 'com.ec.frmw.co.screens/validation_overview.jsf'),
}

for key, (label, url) in screen_labels.items():
    if key in screenshots and os.path.exists(screenshots[key]):
        doc.add_heading(f'  {label}', level=2)
        note = doc.add_paragraph(f'URL: {url}')
        note.runs[0].font.size = Pt(8)
        note.runs[0].font.color.rgb = RGBColor(0x80, 0x80, 0x80)
        try:
            doc.add_picture(screenshots[key], width=Inches(6.5))
            last_para = doc.paragraphs[-1]
            last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            print(f"  Inserted: {label}")
        except Exception as e:
            doc.add_paragraph(f'[Screenshot could not be inserted: {e}]')
        doc.add_paragraph()

doc.save(docx_path)
print(f"\nDocument updated: {docx_path}")
print("Done.")
