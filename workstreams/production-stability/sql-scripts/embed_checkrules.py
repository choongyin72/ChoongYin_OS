"""Embed Check Rule maintenance screenshots: §10.5 frozen rules (insert before §11),
§11.6 sum MOL% rules (append into §11). No commit."""
import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

DOC = r"c:/Projects/ChoongYin_OS/workstreams/production-stability/sql-scripts/Issue1052_Evidence_COPS_DEV.docx"
SS  = r"c:/Projects/ChoongYin_OS/workstreams/master-plan/ec-automation/results/sum_evidence"

def find_heading(doc, prefix):
    for p in doc.paragraphs:
        if p.style.name.startswith('Heading') and p.text.strip().startswith(prefix):
            return p
    return None

def block_before(target, title, cap, path, level='Heading 3'):
    target.insert_paragraph_before(title, style=level)
    c = target.insert_paragraph_before(cap); c.runs[0].font.size = Pt(8.5); c.runs[0].font.color.rgb = RGBColor(0x60,0x60,0x60)
    pic = target.insert_paragraph_before(); pic.add_run().add_picture(path, width=Inches(6.8)); pic.alignment = WD_ALIGN_PARAGRAPH.CENTER
    target.insert_paragraph_before('')

def block_end(doc, title, cap, path, level=3):
    doc.add_heading(title, level=level)
    c = doc.add_paragraph(cap); c.runs[0].font.size = Pt(8.5); c.runs[0].font.color.rgb = RGBColor(0x60,0x60,0x60)
    doc.add_picture(path, width=Inches(6.8)); doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

doc = Document(DOC)
# §10.5 frozen rules — insert before §11
target = find_heading(doc, '11. Sum')
target.insert_paragraph_before('10.5  Frozen rules deployed — Check Rule maintenance screen', style='Heading 2')
ip = target.insert_paragraph_before(
    'Check Rule screen (Configuration > System) filtered to "FROZEN_V1": all 6 frozen rules '
    '(1150-1155) configured — CHECK_ID, name, RV view, Where Formula (${FunctionFrozenValue} = '
    '${ConstBOOLEAN}), Check Message and Severity (Warning). 1150/1151 are the ON-HOLD composition rules.')
ip.runs[0].font.size = Pt(9)
pic = target.insert_paragraph_before(); pic.add_run().add_picture(os.path.join(SS,'check_rule_FROZEN.png'), width=Inches(6.8))
pic.alignment = WD_ALIGN_PARAGRAPH.CENTER
target.insert_paragraph_before('')

# §11.6 sum MOL% rules — append (into §11, which is the last section)
block_end(doc, '11.6  Sum MOL% rules deployed — Check Rule maintenance screen',
    'Check Rule screen filtered to "GAS_COMPONENT_COMP_MOL_PCT": the two new sum rules 1156 '
    '(stream) and 1157 (well) configured — RV view RV_STRM/WELL_GAS_ANALYSIS, Where Formula '
    '(${isComponentSumOutOfTolerance} = ${ConstYES}), Severity ERROR.',
    os.path.join(SS,'check_rule_MOLPCT.png'), level=2)

doc.save(DOC)
print("Embedded §10.5 (frozen rules) + §11.6 (sum MOL rules).")
