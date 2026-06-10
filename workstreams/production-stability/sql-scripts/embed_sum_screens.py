"""Embed the 3 isolated SUM-check screen captures into evidence doc as Section 11.5.
Validation Overview - Pluto Scarborough, Message column filtered to isolate each rule.
Preserves existing content. Does NOT commit."""
import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

DOC = r"c:/Projects/ChoongYin_OS/workstreams/production-stability/sql-scripts/Issue1052_Evidence_COPS_DEV.docx"
SS  = r"c:/Projects/ChoongYin_OS/workstreams/master-plan/ec-automation/results/sum_evidence"

shots = [
    ("sum_stream_MOLE_pct.png",
     "11.5.1  Stream Gas Component Analysis — MOLE % sum check (rule 1156, NEW)",
     "Validation Overview - Pluto Scarborough > Daily Sampling Validations > Stream Gas "
     "Component Analysis group, May 2026. Message column filtered to 'mole percentage': every "
     "row is the new COMP_MOL_PCT sum-check ERROR firing on screen (group status ERROR)."),
    ("sum_stream_WT_pct.png",
     "11.5.2  Stream Gas Component Analysis — WEIGHT % sum check (rule 1077)",
     "Same group/date, Message filtered to 'molecular weight percentage': the existing "
     "COMP_WT_PCT sum-check ERRORs (verified still firing post-fix)."),
    ("sum_well_MOLE_pct.png",
     "11.5.3  Well Gas Component Analysis — MOLE % sum check (rule 1157, NEW)",
     "Well Gas Component Analysis group, 2026-06-01, filtered to 'mole percentage': 13 SCA "
     "wells firing the new COMP_MOL_PCT sum-check ERROR (wells carry no mole % data)."),
]

doc = Document(DOC)
doc.add_heading('11.5  EC Web Screen Evidence — Validation Overview (sum-check ERRORs on screen)', level=2)
doc.add_paragraph(
    'Captured headless via Robot Framework on COPS DEV (sum_check_evidence / sum_one_capture). '
    'The Validation Overview screen runs the daily-sampling gas-component groups; the result grid '
    'is filtered on the Message column to isolate each rule. Row counts were asserted in the run '
    '(mole-only: molecular-weight=0; WT-only: mole=0).'
).runs[0].font.size = Pt(9)

for fn, title, cap in shots:
    path = os.path.join(SS, fn)
    if not os.path.exists(path):
        doc.add_paragraph(f'[missing screenshot: {fn}]')
        continue
    h = doc.add_heading(title, level=3)
    c = doc.add_paragraph(cap)
    c.runs[0].font.size = Pt(8.5); c.runs[0].font.color.rgb = RGBColor(0x60, 0x60, 0x60)
    doc.add_picture(path, width=Inches(6.8))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

doc.save(DOC)
print("Embedded 11.5 with", sum(1 for fn,_,_ in shots if os.path.exists(os.path.join(SS,fn))), "screenshots.")
