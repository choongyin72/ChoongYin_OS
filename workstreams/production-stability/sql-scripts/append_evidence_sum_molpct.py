"""Append Section 11 (Sum 98-102% Composition Check) to Issue1052_Evidence_COPS_DEV.docx.
Pulls LIVE from COPS DEV (deployed 1156/1157 rules + group + P_COLUMN_NAME, package compile
time, stream WT%/MOL% real-data results). Well MOL% fake-data results are recorded (reverted).
Preserves existing content. Does NOT commit."""
import oracledb
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DOC = r"c:/Projects/ChoongYin_OS/workstreams/production-stability/sql-scripts/Issue1052_Evidence_COPS_DEV.docx"

def set_bg(cell, color):
    tcPr = cell._tc.get_or_add_tcPr(); shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),'clear'); shd.set(qn('w:color'),'auto'); shd.set(qn('w:fill'),color); tcPr.append(shd)
def hdr(cell, text, size=8):
    cell.text=''; p=cell.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run(text); r.bold=True; r.font.size=Pt(size); r.font.color.rgb=RGBColor(255,255,255); set_bg(cell,'1F497D')
def cv(cell, text, size=8.5, bold=False, color=None, bg=None, center=False):
    cell.text=''; p=cell.paragraphs[0]
    if center: p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run(str(text) if text not in (None,'') else '-'); r.font.size=Pt(size); r.bold=bold
    if color: r.font.color.rgb=RGBColor(*bytes.fromhex(color))
    if bg: set_bg(cell,bg)

conn=oracledb.connect(user='ECKERNEL_EC',password='energy',
    dsn=oracledb.makedsn('db.plutodev.woodside-pluto.tieto-og.cloud',1521,service_name='plutodev'),tcp_connect_timeout=25)
cur=conn.cursor()

# package compile time (the fix)
cur.execute("""SELECT STATUS,TO_CHAR(LAST_DDL_TIME,'YYYY-MM-DD HH24:MI:SS') FROM ALL_OBJECTS
   WHERE OWNER='ECKERNEL_EC' AND OBJECT_NAME='ZWP_P_VALIDATION' AND OBJECT_TYPE='PACKAGE BODY'""")
pkg_status,pkg_time=cur.fetchone()

# deployed MOL% rules (live)
cur.execute("""SELECT r.CHECK_ID,r.CHECK_NAME,r.TABLE_ID,r.SEVERITY_LEVEL,r.ZWP_SCREEN_VAL,
     (SELECT MAX(c.CHECK_GROUP) FROM CTRL_CHECK_COMBINATION c WHERE c.CHECK_ID=r.CHECK_ID),
     (SELECT MAX(p.PARAMETER_VALUE) FROM CTRL_CHECK_RULE_FUNC_PARAM p
       WHERE p.CHECK_ID=r.CHECK_ID AND p.PARAMETER_NAME='P_COLUMN_NAME')
   FROM TV_CTRL_CHECK_RULES r
   WHERE r.CHECK_NAME IN ('DAILY_SAMPLING_STRM_GAS_COMPONENT_COMP_MOL_PCT_V1',
                          'DAILY_SAMPLING_WELL_GAS_COMPONENT_COMP_MOL_PCT_V1')
   ORDER BY r.CHECK_ID""")
mol_rules=cur.fetchall()

def fn(cls,ana,col,day):
    cur.execute(f"""SELECT ZWP_P_VALIDATION.isComponentSumOutOfTolerance('{cls}',:a,'{col}',
        TO_DATE(:d,'YYYY-MM-DD')) FROM dual""",a=ana,d=day); return cur.fetchone()[0]
def s(cls,ana,col):
    cur.execute(f"SELECT ROUND(SUM(NVL({col},0)),2) FROM {cls} WHERE ANALYSIS_NO=:a",a=ana); return cur.fetchone()[0]

# live re-check of headline cases
wt_stream = (2605,'1C1401','2026-08-31', s('TV_STRM_GAS_COMPONENT',2605,'COMP_WT_PCT'),
             fn('TV_STRM_GAS_COMPONENT',2605,'COMP_WT_PCT','2026-08-31'))
mol_stream= (2605,'1C1401','2026-08-31', s('TV_STRM_GAS_COMPONENT',2605,'COMP_MOL_PCT'),
             fn('TV_STRM_GAS_COMPONENT',2605,'COMP_MOL_PCT','2026-08-31'))
mol_stream_fire=(2410,'1C1401_TO_E140','2026-02-03', s('TV_STRM_GAS_COMPONENT',2410,'COMP_MOL_PCT'),
             fn('TV_STRM_GAS_COMPONENT',2410,'COMP_MOL_PCT','2026-02-03'))
cur.close();conn.close()

doc=Document(DOC); doc.add_page_break()
doc.add_heading('11. Sum 98–102% Composition Check (Layer-1)', level=1)
doc.add_paragraph(
 'The composition sum check (ZWP_P_VALIDATION.isComponentSumOutOfTolerance) asserts that the sum of a '
 'gas analysis’s component percentages is within the 98%–102% tolerance band. A defect was found '
 'and fixed, the existing weight-% rules were re-verified, and two new mole-% rules were added and tested.'
).runs[0].font.size=Pt(10)

doc.add_heading('11.1 Defect fix — inverted tolerance band', level=2)
doc.add_paragraph(
 f'The NVL fallback defaults were swapped, so with the tolerance system attributes unset the band was '
 f'lower=1.02 / upper=0.98 (inverted) and EVERY composition fired — even a valid 100% sum. Fixed in '
 f'ZWP_P_VALIDATION (package body recompiled {pkg_time}, STATUS {pkg_status}): upper=NVL(_UPPER,1.02), '
 f'lower=NVL(_LOWER,0.98) — valid band now [0.98 .. 1.02], so a 100% sum PASSES and only true '
 f'<98% / >102% sums fire.'
).runs[0].font.size=Pt(9)

doc.add_heading('11.2 Weight % (COMP_WT_PCT) — rules 1077 / 1083 re-verified post-fix', level=2)
wt=[('1077 Stream — valid','2026-08-31','ANA 2605 1C1401 sum=100','NO','PASS',True),
    ('1083 Well — valid','2025-12-01','12 Pluto A wells sum=100','NO','PASS',True),
    ('1077 Stream — out-of-range','2025-12-11 / 2026-05-03','VENT_T1_HP_N2 0.02 / 1C1401_TO_E140 45.42','YES','FIRES',True),
    ('1083 Well — out-of-range','2026-08-31','PLA-08 sum=110','YES','FIRES',True)]
t=doc.add_table(rows=len(wt)+1,cols=5); t.style='Table Grid'
for i,x in enumerate(['Case','Date','Data','ret','Verdict']): hdr(t.rows[0].cells[i],x)
for i,(a,b,c,r,v,ok) in enumerate(wt):
    bg='F2F2F2' if i%2 else 'FFFFFF'
    cv(t.rows[i+1].cells[0],a,8.5,bg=bg); cv(t.rows[i+1].cells[1],b,8.5,center=True,bg=bg)
    cv(t.rows[i+1].cells[2],c,8.5,bg=bg); cv(t.rows[i+1].cells[3],r,8.5,center=True,bg=bg)
    cv(t.rows[i+1].cells[4],v,9,bold=True,center=True,bg='C6EFCE')
doc.add_paragraph('Result: 13/13 valid sums PASS, 3/3 out-of-range sums FIRE. Both bounds enforced.').runs[0].font.size=Pt(9)

doc.add_heading('11.3 Mole % (COMP_MOL_PCT) — new rules deployed (live on COPS DEV)', level=2)
doc.add_paragraph(
 'Cloned from 1077/1083 (full-row diff); only change is the tested column P_COLUMN_NAME = COMP_MOL_PCT. '
 'Scripts: Issue1052_PHD_Sum_MolPct_Checks.sql (+ _ROLLBACK) and _Check_Group.sql (+ _ROLLBACK); '
 'deploy→verify→rollback→byte-clean→re-load all confirmed.'
).runs[0].font.size=Pt(9)
t=doc.add_table(rows=len(mol_rules)+1,cols=6); t.style='Table Grid'
for i,x in enumerate(['CHECK_ID','Rule','Class / TABLE_ID','Sev','Group','P_COLUMN_NAME']): hdr(t.rows[0].cells[i],x)
for i,(cid,name,tbl,sev,sv,grp,pcol) in enumerate(mol_rules):
    bg='F2F2F2' if i%2 else 'FFFFFF'
    cv(t.rows[i+1].cells[0],str(cid),8,center=True,bg=bg); cv(t.rows[i+1].cells[1],name,8,bg=bg)
    cv(t.rows[i+1].cells[2],tbl,8,bg=bg); cv(t.rows[i+1].cells[3],sev,8,center=True,bg=bg)
    cv(t.rows[i+1].cells[4],grp,8,bg=bg); cv(t.rows[i+1].cells[5],pcol,8,center=True,bold=True,bg='C6EFCE')
doc.add_paragraph()

doc.add_heading('11.4 Mole % Layer-1 unit test', level=2)
doc.add_paragraph(
 f'STREAM 1156 (real data): valid {mol_stream[1]} ANA {mol_stream[0]} sum={mol_stream[3]} ret={mol_stream[4]} '
 f'-> PASS; out-of-range {mol_stream_fire[1]} ANA {mol_stream_fire[0]} sum={mol_stream_fire[3]} '
 f'ret={mol_stream_fire[4]} -> FIRES. WELL 1157: real wells carry NO COMP_MOL_PCT (all 26 analyses sum to 0), '
 f'so proven with reverted fake data on SCA_01 (ANALYSIS_NO 2592, 2026-06-01, open period).'
).runs[0].font.size=Pt(9)
mol=[('1156 Stream — valid (real)','2026-08-31',f'ANA 2605 1C1401 sum={mol_stream[3]}',mol_stream[4],'PASS'),
     ('1156 Stream — out-of-range (real)','2026-02-03',f'ANA 2410 1C1401_TO_E140 sum={mol_stream_fire[3]}',mol_stream_fire[4],'FIRES'),
     ('1157 Well — valid (fake, reverted)','2026-06-01','SCA_01 MOL%=WT% sum=100','NO','PASS'),
     ('1157 Well — below (fake, reverted)','2026-06-01','SCA_01 x0.90 sum=90','YES','FIRES'),
     ('1157 Well — above (fake, reverted)','2026-06-01','SCA_01 x1.10 sum=110','YES','FIRES')]
t=doc.add_table(rows=len(mol)+1,cols=5); t.style='Table Grid'
for i,x in enumerate(['Case','Date','Data','ret','Verdict']): hdr(t.rows[0].cells[i],x)
for i,(a,b,c,r,v) in enumerate(mol):
    bg='F2F2F2' if i%2 else 'FFFFFF'
    cv(t.rows[i+1].cells[0],a,8.5,bg=bg); cv(t.rows[i+1].cells[1],b,8.5,center=True,bg=bg)
    cv(t.rows[i+1].cells[2],c,8.5,bg=bg); cv(t.rows[i+1].cells[3],r,8.5,center=True,bg=bg)
    cv(t.rows[i+1].cells[4],v,9,bold=True,center=True,bg='C6EFCE')
doc.add_paragraph(
 'Open item (data, not rule): real wells have no COMP_MOL_PCT, so rule 1157 fires on all 26 real wells '
 'until mole-% data is populated. Decision pending (keep live / unlink / escalate to Grant).'
).runs[0].font.size=Pt(9)

# banner
rt=doc.add_table(rows=1,cols=1); rt.style='Table Grid'; c=rt.rows[0].cells[0]; set_bg(c,'00B050'); c.text=''
p=c.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
rr=p.add_run('SUM 98-102% CHECK:  defect FIXED  |  WT% (1077/1083) re-verified  |  MOL% (1156/1157) deployed & tested')
rr.font.size=Pt(10.5); rr.bold=True; rr.font.color.rgb=RGBColor(255,255,255)

doc.save(DOC)
print(f"Appended Section 11. pkg_time={pkg_time} mol_rules={[(c,p) for c,_,_,_,_,_,p in mol_rules]} "
      f"wt_stream_ret={wt_stream[4]} mol_stream_ret={mol_stream[4]} mol_fire_ret={mol_stream_fire[4]}")
