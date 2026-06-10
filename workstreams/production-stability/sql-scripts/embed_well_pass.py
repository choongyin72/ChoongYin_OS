import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
DOC=r"c:/Projects/ChoongYin_OS/workstreams/production-stability/sql-scripts/Issue1052_Evidence_COPS_DEV.docx"
SS=r"c:/Projects/ChoongYin_OS/workstreams/master-plan/ec-automation/results/sum_evidence"
doc=Document(DOC)
doc.add_heading('11.7  Well MOL% rule PASSES with valid data — fake-data test (reverted)', level=2)
c=doc.add_paragraph(
 'Real wells carry no COMP_MOL_PCT (all fire). To prove rule 1157 PASSES valid data (not just '
 'fires on missing data), SCA_01 was temporarily patched to a valid 100% mole sum '
 '(Issue1052_PHD_Sum_MolPct_FakeData_Patch.sql), the well group re-run, then reverted. The mole% '
 'errors drop 13 -> 12: SCA 02-13 still fire, SCA 01 is absent (PASSES). COPS DEV reverted '
 'byte-clean (SCA_01 MOL% back to NULL; CTRL_CHECK_LOG restored to 13).')
c.runs[0].font.size=Pt(8.5); c.runs[0].font.color.rgb=RGBColor(0x60,0x60,0x60)
doc.add_picture(os.path.join(SS,'sum_well_PASS_fakedata.png'), width=Inches(6.8))
doc.paragraphs[-1].alignment=WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph()
doc.save(DOC)
print("Embedded 11.7 (well fake-data PASS).")
