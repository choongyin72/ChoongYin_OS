"""Insert Phase-1 null/range screen evidence as Section 5.5, before Section 6. No commit."""
import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

DOC = r"c:/Projects/ChoongYin_OS/workstreams/production-stability/sql-scripts/Issue1052_Evidence_COPS_DEV.docx"
SS  = r"c:/Projects/ChoongYin_OS/workstreams/master-plan/ec-automation/results/sum_evidence"

def find_heading(doc, prefix):
    for p in doc.paragraphs:
        if p.style.name.startswith('Heading') and p.text.strip().startswith(prefix):
            return p
    return None

doc = Document(DOC)
target = find_heading(doc, '6.')
assert target is not None, "no §6 heading"

blocks = [
 ('5.5.1  Stream Gas Component Analysis (Composition) — Mol%/Wt% null-range (1142/1143)',
  'Validation Overview > Stream Gas Component Analysis (Composition) - PHD Validations, '
  '2026-05-26, Message filtered to "invalid or missing": the MOL_PCT (TC01) and WT_PCT (TC02) '
  'null/range ERRORs firing on screen.',
  os.path.join(SS, 'phase1_composition_2026-05-26.png')),
 ('5.5.2  Stream Gas Component Analysis (Analysis) — Density/GCV null-range (1144/1145)',
  'Stream Gas Component Analysis (Analysis) - PHD Validations, 2026-05-26: DENSITY (TC03) and '
  'GCV (TC04) null/range ERRORs.',
  os.path.join(SS, 'phase1_analysis_2026-05-26.png')),
 ('5.5.3  Daily Tank Status — tank null-range (1146–1149)',
  'Daily Tank Status - VCF Calc - PHD Validations, 2026-06-07: Gross Volume/Mass, Avg Temp and '
  'Std Density null/range ERRORs (TC05–TC08).',
  os.path.join(SS, 'phase1_tank_2026-06-07.png')),
]

h = target.insert_paragraph_before('5.5  EC Web Screen Evidence — Phase-1 null/range ERRORs (Validation Overview)', style='Heading 2')
ip = target.insert_paragraph_before(
    'Headless RF capture on COPS DEV. Each PHD validation group is selected on the Validation '
    'Overview screen and the result grid filtered on the Message column to isolate the null/range '
    'ERRORs of the Phase-1 rules (1142-1149).'); ip.runs[0].font.size = Pt(9)
for title, cap, path in blocks:
    target.insert_paragraph_before(title, style='Heading 3')
    c = target.insert_paragraph_before(cap); c.runs[0].font.size = Pt(8.5); c.runs[0].font.color.rgb = RGBColor(0x60,0x60,0x60)
    pic = target.insert_paragraph_before(); pic.add_run().add_picture(path, width=Inches(6.8)); pic.alignment = WD_ALIGN_PARAGRAPH.CENTER
    target.insert_paragraph_before('')

doc.save(DOC)
print("Inserted 5.5 (Phase-1 screen evidence) before §6.")
